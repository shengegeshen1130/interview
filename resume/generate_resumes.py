#!/usr/bin/env python3
"""Generate English and Chinese one-page AI Engineer resumes for Shen Ge."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ── Shared styling helpers ──────────────────────────────────────────────

def set_narrow_margins(doc):
    """Set narrow margins for one-page fit."""
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

def set_paragraph_spacing(p, before=0, after=0, line_spacing=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)

def add_thin_border(paragraph):
    """Add a thin bottom border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="333333"/>'
        '</w:pBdr>'
    )
    pPr.append(borders)

def add_run(paragraph, text, bold=False, italic=False, size=None, color=None, font_name=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name
    return run

def set_default_font(doc, name="Calibri", size=10, east_asia=None):
    style = doc.styles['Normal']
    font = style.font
    font.name = name
    font.size = Pt(size)
    if east_asia:
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), east_asia)

# ── English Resume ──────────────────────────────────────────────────────

def create_english_resume():
    doc = Document()
    set_narrow_margins(doc)
    set_default_font(doc, "Calibri", 10)

    # ── Name ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=22)
    add_run(p, "SHEN GE", bold=True, size=20, color=(0x1A, 0x1A, 0x1A))

    # ── Contact ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=2, line_spacing=13)
    add_run(p, "1339105051@qq.com  |  +86-186-0205-4952  |  Shenzhen, China", size=9, color=(0x44, 0x44, 0x44))

    # ── Summary ──
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=2, line_spacing=13)
    add_thin_border(p)
    add_run(p, "PROFESSIONAL SUMMARY", bold=True, size=11, color=(0x1A, 0x1A, 0x1A))

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=3, line_spacing=13)
    add_run(p, (
        "Senior ML/AI Engineer with 6+ years at PayPal building production ML and LLM systems for fraud detection, "
        "anti-money laundering, and financial risk. Deep expertise in multi-agent LLM orchestration (LangChain, LangGraph), "
        "Graph ML (GNN, community detection), and end-to-end model lifecycle from problem framing to deployment at scale. "
        "Proven track record of delivering high-impact AI solutions — $100M+ annual fraud loss savings, 85% analyst efficiency gain, "
        "and 94% SQL generation accuracy. Seeking AI Engineer roles to apply production LLM and ML expertise to next-generation intelligent systems."
    ), size=9.5)

    # ── Skills ──
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=2, line_spacing=13)
    add_thin_border(p)
    add_run(p, "TECHNICAL SKILLS", bold=True, size=11, color=(0x1A, 0x1A, 0x1A))

    skills = [
        ("LLM & Agents", "LangChain, LangGraph, DSPy, Prompt Engineering, RAG, Graph-RAG, Multi-Agent Systems, Chain-of-Thought"),
        ("ML & Deep Learning", "PyTorch, DGL, GNN (GAT), AutoEncoder, LightGBM, XGBoost, PU-Learning, Anomaly Detection"),
        ("Data & Infrastructure", "BigQuery, Spark, Neo4j, Gremlin, Pandas, SQL, Airflow, Docker, GCP, Model Serving"),
        ("Languages & Tools", "Python, Scala, Java, Git, Jupyter, SHAP, PandasAI, Scikit-learn"),
    ]
    for label, items in skills:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0, line_spacing=13)
        add_run(p, f"{label}: ", bold=True, size=9)
        add_run(p, items, size=9)

    # ── Experience ──
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=2, line_spacing=13)
    add_thin_border(p)
    add_run(p, "PROFESSIONAL EXPERIENCE", bold=True, size=11, color=(0x1A, 0x1A, 0x1A))

    # PayPal header
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=1, line_spacing=13)
    add_run(p, "PayPal", bold=True, size=10.5)
    add_run(p, "  —  Senior Machine Learning Engineer", bold=False, size=10)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p2, before=0, after=2, line_spacing=13)
    add_run(p2, "Shenzhen, China  |  Feb 2019 – Present", italic=True, size=9, color=(0x55, 0x55, 0x55))

    # ── Project 1: AML Investigation Mate ──
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=3, after=1, line_spacing=13)
    add_run(p, "AML Investigation Mate — Multi-Agent LLM System", bold=True, size=9.5)
    add_run(p, "  (Jan 2025 – Present)", italic=True, size=9, color=(0x55, 0x55, 0x55))

    bullets_aml = [
        "Architected a multi-agent system (LangChain + deep-agent pattern) automating end-to-end AML case investigation — evidence retrieval, regulatory rule reasoning, risk scoring, and SAR report generation — achieving 80% decision accuracy on par with senior investigators.",
        "Built a Graph-RAG knowledge layer (Neo4j) encoding 50+ internal SOPs and external AML regulations into a structured knowledge graph with condition-action-legal linkages, enabling grounded agent reasoning over complex compliance logic.",
        "Reduced average case review time from 3–4 hours to under 30 minutes (85% reduction) across 15 investigators through automatic prompt optimization (DSPy) and iterative multi-agent evaluation pipelines.",
    ]
    for b in bullets_aml:
        p = doc.add_paragraph(style='List Bullet')
        set_paragraph_spacing(p, before=0, after=0, line_spacing=12.5)
        p.clear()
        add_run(p, b, size=9)

    # ── Project 2: Text-to-SQL ──
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=3, after=1, line_spacing=13)
    add_run(p, "Text-to-SQL Investigation Toolkit — LLM-Powered Query Generation", bold=True, size=9.5)
    add_run(p, "  (Jan 2025 – Present)", italic=True, size=9, color=(0x55, 0x55, 0x55))

    bullets_sql = [
        "Designed a multi-step RAG pipeline (BGE-M3 embeddings + Chain-of-Thought prompting) converting natural language to BigQuery SQL with 94% correctness on 200-query golden dataset, reducing manual SQL coding time by 80%.",
        "Built a stateful multi-agent workflow (LangGraph) integrating intent routing, error self-healing (3-retry dry-run feedback loop), interactive SQL refinement, data visualization (PandasAI), and feedback reuse in a unified human-in-the-loop platform.",
    ]
    for b in bullets_sql:
        p = doc.add_paragraph(style='List Bullet')
        set_paragraph_spacing(p, before=0, after=0, line_spacing=12.5)
        p.clear()
        add_run(p, b, size=9)

    # ── Project 3: Graph Fraud Detection ──
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=3, after=1, line_spacing=13)
    add_run(p, "Proactive Fraud Detection — Graph-Based Account Linking & Clustering", bold=True, size=9.5)
    add_run(p, "  (Jun 2023 – Dec 2024)", italic=True, size=9, color=(0x55, 0x55, 0x55))

    bullets_graph = [
        "Led end-to-end development of a graph fraud detection pipeline processing 100M+ accounts and 300M edges, combining seed-based linking, Louvain community detection, and embedding similarity to surface fraud rings early in account lifecycle — delivering $100M+ annual net loss savings.",
        "Trained AutoEncoder for account embedding extraction via model distillation, then applied edge-aware GAT (DGL/PyTorch) on the linking graph to compute group-level and account-level risk scores; formulated investigator allocation as ILP for optimal resource assignment.",
    ]
    for b in bullets_graph:
        p = doc.add_paragraph(style='List Bullet')
        set_paragraph_spacing(p, before=0, after=0, line_spacing=12.5)
        p.clear()
        add_run(p, b, size=9)

    # ── Project 4: ML Model Dev ──
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=3, after=1, line_spacing=13)
    add_run(p, "ML Model Development — Multi-Domain Risk Modeling", bold=True, size=9.5)
    add_run(p, "  (Feb 2019 – May 2023)", italic=True, size=9, color=(0x55, 0x55, 0x55))

    bullets_ml = [
        "Owned full model lifecycle across 8+ risk domains (stolen finance, account takeover, buyer AUP violation, collusion, merchant website compliance, dispute automation) using LightGBM, DNN, PU-Learning, and sequence-based anomaly detection.",
        "Pioneered Word2Vec buyer embeddings and KMeans behavioral clustering for AUP violation detection; designed loss-weighted DNN training for stolen financial instruments; delivered explainable AI outputs (SHAP) for stakeholder trust and regulatory compliance.",
    ]
    for b in bullets_ml:
        p = doc.add_paragraph(style='List Bullet')
        set_paragraph_spacing(p, before=0, after=0, line_spacing=12.5)
        p.clear()
        add_run(p, b, size=9)

    # Leadership
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=3, after=1, line_spacing=13)
    add_run(p, "Leadership: ", bold=True, size=9)
    add_run(p, "Mentored 2 junior data scientists on ML development and fraud domain expertise. Led PayPal's Recommendation System Study Group (2024), driving cross-team knowledge sharing on state-of-the-art RecSys.", size=9)

    # ── Education ──
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=2, line_spacing=13)
    add_thin_border(p)
    add_run(p, "EDUCATION", bold=True, size=11, color=(0x1A, 0x1A, 0x1A))

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=0, line_spacing=13)
    add_run(p, "The Chinese University of Hong Kong", bold=True, size=9.5)
    add_run(p, "  —  MSc in Business Analytics  |  Sep 2017 – Nov 2018", size=9)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=0, line_spacing=13)
    add_run(p, "College of William & Mary", bold=True, size=9.5)
    add_run(p, "  —  B.S. in Applied Mathematics (CS Minor)  |  Aug 2013 – May 2017  |  GPA: 3.62/4.0, 6x Dean's List", size=9)

    path = os.path.join(OUTPUT_DIR, "ShenGE_Resume_AI_Engineer_EN.docx")
    doc.save(path)
    print(f"English resume saved: {path}")
    return path


# ── Chinese Resume ──────────────────────────────────────────────────────

def create_chinese_resume():
    doc = Document()
    set_narrow_margins(doc)
    set_default_font(doc, "Calibri", 10, east_asia="Microsoft YaHei")

    # ── Name ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=22)
    r = add_run(p, "葛  深", bold=True, size=20, color=(0x1A, 0x1A, 0x1A))
    r.font.name = "Microsoft YaHei"
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), "Microsoft YaHei")

    # ── Contact ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=2, line_spacing=13)
    add_run(p, "1339105051@qq.com  |  +86-186-0205-4952  |  深圳", size=9, color=(0x44, 0x44, 0x44))

    # ── Summary ──
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=2, line_spacing=13)
    add_thin_border(p)
    add_run(p, "个人简介", bold=True, size=11, color=(0x1A, 0x1A, 0x1A))

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=3, line_spacing=13.5)
    add_run(p, (
        "高级 ML/AI 工程师，6年以上 PayPal 生产级 ML 及 LLM 系统设计与落地经验。"
        "擅长 Multi-Agent LLM 编排（LangChain、LangGraph）、Graph ML（GNN、Community Detection）及模型全生命周期管理。"
        "主导交付多个高影响力 AI 项目——年化 $100M+ Fraud Loss Savings、分析师效率提升 85%、SQL 生成准确率 94%。"
        "现寻求 AI Engineer 岗位，将 Production LLM 与 ML 工程经验应用于下一代智能系统。"
    ), size=9.5)

    # ── Skills ──
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=2, line_spacing=13)
    add_thin_border(p)
    add_run(p, "技术技能", bold=True, size=11, color=(0x1A, 0x1A, 0x1A))

    skills = [
        ("LLM & Agents", "LangChain, LangGraph, DSPy, Prompt Engineering, RAG, Graph-RAG, Multi-Agent Systems, Chain-of-Thought"),
        ("ML & Deep Learning", "PyTorch, DGL, GNN (GAT), AutoEncoder, LightGBM, XGBoost, PU-Learning, Anomaly Detection"),
        ("数据与基础设施", "BigQuery, Spark, Neo4j, Gremlin, Pandas, SQL, Airflow, Docker, GCP, Model Serving"),
        ("编程语言与工具", "Python, Scala, Java, Git, Jupyter, SHAP, PandasAI, Scikit-learn"),
    ]
    for label, items in skills:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0, line_spacing=13)
        add_run(p, f"{label}: ", bold=True, size=9)
        add_run(p, items, size=9)

    # ── Experience ──
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=2, line_spacing=13)
    add_thin_border(p)
    add_run(p, "工作经历", bold=True, size=11, color=(0x1A, 0x1A, 0x1A))

    # PayPal header
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=1, line_spacing=13)
    add_run(p, "PayPal", bold=True, size=10.5)
    add_run(p, "  —  高级机器学习工程师 (Senior Machine Learning Engineer)", size=10)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p2, before=0, after=2, line_spacing=13)
    add_run(p2, "深圳  |  2019年2月 – 至今", italic=True, size=9, color=(0x55, 0x55, 0x55))

    # ── Project 1: AML ──
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=3, after=1, line_spacing=13)
    add_run(p, "AML Investigation Mate — Multi-Agent LLM 智能调查系统", bold=True, size=9.5)
    add_run(p, "  (2025.01 – 至今)", italic=True, size=9, color=(0x55, 0x55, 0x55))

    bullets_aml = [
        "基于 LangChain Deep Agent 搭建 Multi-Agent 架构，实现 AML 案件端到端自动调查——涵盖 Evidence Retrieval、Regulatory Rule Reasoning、Risk Scoring 和 SAR Report 生成——Decision Accuracy 达 80%，与资深调查员持平。",
        "构建 Graph-RAG 知识层（Neo4j），将 50+ 内部 SOP 和外部 AML 法规编码为结构化知识图谱（Condition → Action → Legal Document），赋予 Agent 合规推理能力。",
        "通过 DSPy 自动 Prompt 优化和迭代评估，将单案审查时间从 3–4 小时缩短至 30 分钟以内（效率提升 85%），覆盖 15 名调查员的实测验证。",
    ]
    for b in bullets_aml:
        p = doc.add_paragraph(style='List Bullet')
        set_paragraph_spacing(p, before=0, after=0, line_spacing=12.5)
        p.clear()
        add_run(p, b, size=9)

    # ── Project 2: Text-to-SQL ──
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=3, after=1, line_spacing=13)
    add_run(p, "Text-to-SQL 调查工具 — LLM 驱动的自然语言查询系统", bold=True, size=9.5)
    add_run(p, "  (2025.01 – 至今)", italic=True, size=9, color=(0x55, 0x55, 0x55))

    bullets_sql = [
        "设计 Multi-step RAG Pipeline（BGE-M3 Embedding + Chain-of-Thought Prompting），将自然语言转为 BigQuery SQL，在 200 条 Golden Dataset 上达 94% Correctness，人工编码时间减少 80%。",
        "基于 LangGraph 构建 Stateful Multi-Agent Workflow，集成 Intent Routing、Error Self-Healing（3 轮 Dry Run Feedback Loop）、交互式 SQL 修改、PandasAI 可视化分析和 Feedback 复用，形成统一 Human-in-the-Loop 调查平台。",
    ]
    for b in bullets_sql:
        p = doc.add_paragraph(style='List Bullet')
        set_paragraph_spacing(p, before=0, after=0, line_spacing=12.5)
        p.clear()
        add_run(p, b, size=9)

    # ── Project 3: Graph Fraud ──
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=3, after=1, line_spacing=13)
    add_run(p, "主动欺诈检测 — Graph-Based Account Linking & Clustering", bold=True, size=9.5)
    add_run(p, "  (2023.06 – 2024.12)", italic=True, size=9, color=(0x55, 0x55, 0x55))

    bullets_graph = [
        "主导开发基于图的 Fraud Detection Pipeline，处理 1 亿+ 账户和 3 亿条边，结合 Seed-Based Linking、Louvain Community Detection 和 Embedding Similarity 在账户早期识别欺诈团伙——年化 Net Loss Savings 超 $100M。",
        "通过 Model Distillation 训练 AutoEncoder 提取 Account Embedding，结合 Edge-Aware GAT（DGL/PyTorch）计算 Group 和 Account 级别 Risk Score；将调查员资源分配建模为 ILP 优化问题。",
    ]
    for b in bullets_graph:
        p = doc.add_paragraph(style='List Bullet')
        set_paragraph_spacing(p, before=0, after=0, line_spacing=12.5)
        p.clear()
        add_run(p, b, size=9)

    # ── Project 4: ML Model Dev ──
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=3, after=1, line_spacing=13)
    add_run(p, "多风控领域 ML 模型开发 — Multi-Domain Risk Modeling", bold=True, size=9.5)
    add_run(p, "  (2019.02 – 2023.05)", italic=True, size=9, color=(0x55, 0x55, 0x55))

    bullets_ml = [
        "负责 8+ 风控场景的模型全生命周期（Stolen Finance、Account Takeover、Buyer AUP Violation、Collusion、Merchant Website Compliance、Dispute Automation 等），使用 LightGBM、DNN、PU-Learning 和序列异常检测。",
        "创新性地将 Word2Vec Buyer Embedding + KMeans 行为聚类应用于 AUP Violation 检测；设计 Loss-Weighted DNN 训练方案用于盗刷检测；交付 Explainable AI（SHAP）输出满足合规审查要求。",
    ]
    for b in bullets_ml:
        p = doc.add_paragraph(style='List Bullet')
        set_paragraph_spacing(p, before=0, after=0, line_spacing=12.5)
        p.clear()
        add_run(p, b, size=9)

    # Leadership
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=3, after=1, line_spacing=13)
    add_run(p, "团队管理: ", bold=True, size=9)
    add_run(p, "指导 2 名初级数据科学家，帮助其快速上手 ML 模型开发和风控业务。组织 PayPal 推荐系统学习小组（2024），推动跨团队 RecSys 技术分享。", size=9)

    # ── Education ──
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=2, line_spacing=13)
    add_thin_border(p)
    add_run(p, "教育背景", bold=True, size=11, color=(0x1A, 0x1A, 0x1A))

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=0, line_spacing=13)
    add_run(p, "香港中文大学 (The Chinese University of Hong Kong)", bold=True, size=9.5)
    add_run(p, "  —  商业分析硕士 (MSc in Business Analytics)  |  2017.09 – 2018.11", size=9)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=0, line_spacing=13)
    add_run(p, "威廉与玛丽学院 (College of William & Mary)", bold=True, size=9.5)
    add_run(p, "  —  应用数学学士 (B.S. Applied Mathematics, CS Minor)  |  2013.08 – 2017.05  |  GPA: 3.62/4.0, 6x Dean's List", size=9)

    path = os.path.join(OUTPUT_DIR, "ShenGE_Resume_AI_Engineer_ZH.docx")
    doc.save(path)
    print(f"Chinese resume saved: {path}")
    return path


if __name__ == "__main__":
    create_english_resume()
    create_chinese_resume()
    print("\nDone! Both resumes generated.")
